import streamlit as st
import streamlit_authenticator as stauth
import sqlite3
import pandas as pd
from streamlit_player import st_player
from docx import Document
import io, base64

# --------------------------------------------------
# PAGE CONFIG (MUST BE FIRST)
# --------------------------------------------------
st.set_page_config(layout="wide", page_title="Video Notebook Manager")

# --------------------------------------------------
# AUTH SETUP (USERNAME / PASSWORD)
# --------------------------------------------------
credentials = {
    "usernames": {
        "divyam": {   # 👈 username
            "name": "Divyam Kumar Pandey",
            "email": "divyamkumarp@gmail.com",
            "password": "$2b$12$FMqWTn51IvhYWYvYZPSCWuZi01wH3u0GF4eBlzbwLOn2cBxAItG6e",  # 👈 bcrypt hash
        }
    }
}

authenticator = stauth.Authenticate(
    credentials,
    cookie_name="youtube_notebook_manager",
    cookie_key=st.secrets["cookie_key"],
    cookie_expiry_days=7,
)

authenticator.login(location="main")

if not st.session_state.get("authentication_status"):
    st.stop()

user_name = st.session_state["name"]
user_email = credentials["usernames"][st.session_state["username"]]["email"]

# --------------------------------------------------
# DATABASE
# --------------------------------------------------
DB_FILE = "notebooks.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS notebooks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_email TEXT NOT NULL,
            title TEXT NOT NULL,
            video_url TEXT,
            notes TEXT,
            progress_time_seconds INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

init_db()

def get_all_notebooks():
    conn = sqlite3.connect(DB_FILE)
    df = pd.read_sql_query(
        "SELECT * FROM notebooks WHERE user_email = ? ORDER BY created_at DESC",
        conn,
        params=(user_email,)
    )
    conn.close()
    return df

def create_notebook(title, url):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        INSERT INTO notebooks (user_email, title, video_url, notes)
        VALUES (?, ?, ?, '')
    """, (user_email, title, url))
    conn.commit()
    conn.close()

def update_notes(nid, notes, progress):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        UPDATE notebooks
        SET notes = ?, progress_time_seconds = ?
        WHERE id = ? AND user_email = ?
    """, (notes, progress, nid, user_email))
    conn.commit()
    conn.close()

def delete_notebook(nid):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute(
        "DELETE FROM notebooks WHERE id = ? AND user_email = ?",
        (nid, user_email)
    )
    conn.commit()
    conn.close()

# --------------------------------------------------
# EXPORT
# --------------------------------------------------
def export_notebook(data):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading(data["title"], level=1)
    doc.add_paragraph(f"Video URL: {data['video_url']}\n\n{data['notes']}")
    doc.save(buffer)
    buffer.seek(0)

    b64 = base64.b64encode(buffer.read()).decode()
    fname = data["title"].replace(" ", "_") + ".docx"

    st.components.v1.html(f"""
    <a download="{fname}"
       href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}">
    </a>
    <script>document.querySelector('a').click()</script>
    """, height=0)

# --------------------------------------------------
# SIDEBAR
# --------------------------------------------------
with st.sidebar:
    st.markdown(f"👋 **{user_name}**")
    authenticator.logout("Logout", "sidebar")
    st.divider()

    mode = st.radio("Menu", ["Open Notebook", "Create New"], label_visibility="collapsed")

    df = get_all_notebooks()
    selected_id = None
    if not df.empty:
        mapping = dict(zip(df["title"], df["id"]))
        selected_title = st.selectbox("Notebooks", mapping.keys())
        selected_id = mapping[selected_title]

# --------------------------------------------------
# MAIN UI
# --------------------------------------------------
if mode == "Create New":
    st.header("✨ Create New Notebook")
    with st.form("create"):
        title = st.text_input("Notebook title")
        url = st.text_input("YouTube URL")
        if st.form_submit_button("Create") and title and url:
            create_notebook(title, url)
            st.success("Notebook created")
            st.rerun()

elif mode == "Open Notebook" and selected_id:
    conn = sqlite3.connect(DB_FILE)
    data = pd.read_sql_query(
        "SELECT * FROM notebooks WHERE id = ?",
        conn,
        params=(selected_id,)
    ).iloc[0]
    conn.close()

    h1, h2, h3 = st.columns([6, 1, 1])
    h1.title(data["title"])
    if h2.button("Export"):
        export_notebook(data)
    if h3.button("Delete"):
        delete_notebook(selected_id)
        st.rerun()

    left, right = st.columns([2, 1])

    with left:
        progressTimeSeconds = int(data['progress_time_seconds'])

        options = {
            "events": ["onProgress"],
            "progress_interval": 500,
            "height": 600,
            "playback_rate": 1.5,
            "config": {
                "youtube": {
                    "playerVars": {
                        "start": progressTimeSeconds
                    }
                }
            }
        }
        
        event = st_player(data['video_url'], **options, key="youtube_player",)
        playedSeconds = 0
        if event :
            (name, data) = event
            playedSeconds = (data or {}).get("playedSeconds", 0)

    with right:
        print(data["notes"])
        notes = st.text_area(
            "Notes",
            value=data['notes'],
            height=500,
            key=f"notes_{selected_id}"
        )
        if notes != data["notes"]:
            update_notes(selected_id, notes, playedSeconds)

else:
    st.info("Create or open a notebook 👈")
