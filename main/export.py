import base64
import io
from pathlib import Path
from textwrap import dedent

from docx import Document
from htmldocx import HtmlToDocx
import streamlit as st
from weasyprint import HTML

ASSETS_PATH = Path(__file__).parent / "assets"


def _load_svg(filename: str) -> str:
    """Load an SVG from the `main/` folder and strip any XML declaration."""
    svg_path = ASSETS_PATH / filename
    raw = Path(svg_path).read_text(encoding="utf-8")
    start = raw.find("<svg")
    return raw[start:] if start != -1 else raw


DOCX_SVG = _load_svg("docx_icon.svg")
PDF_SVG = _load_svg("PDF_file_icon.svg")


def _build_html(notebook_data) -> str:
    """Build a simple HTML page for PDF export."""
    title = notebook_data["title"]
    video_url = notebook_data["video_url"] or ""
    notes_html = notebook_data["notes"] or ""

    html = f"""
    <html>
      <head>
        <meta charset="utf-8">
        <title>{title}</title>
        <style>
          body {{ font-family: sans-serif; margin: 2rem; }}
          h1 {{ margin-bottom: 0.5rem; }}
          .meta {{ color: #555; font-size: 0.9rem; margin-bottom: 1.5rem; }}
          .notes h1, .notes h2, .notes h3 {{ margin-top: 1.5rem; }}
        </style>
      </head>
      <body>
        <h1>{title}</h1>
        <div class="meta">
          Video URL: {video_url}
        </div>
        <div class="notes">
          {notes_html}
        </div>
      </body>
    </html>
    """
    return dedent(html)


@st.dialog("Export notebook")
def _export_dialog(notebook_data) -> None:
    """Dialog UI to choose export format (DOCX or PDF) with large icon cards."""

    # Build DOCX bytes
    docx_buffer = io.BytesIO()
    document = Document()
    document.add_heading(notebook_data["title"], level=1)

    video_url = notebook_data["video_url"]
    if video_url:
        document.add_paragraph(f"Video URL: {video_url}")
        document.add_paragraph("")  # blank line

    notes_html = notebook_data["notes"] or ""
    if notes_html:
        document.add_paragraph("Notes:")
        parser = HtmlToDocx()
        parser.add_html_to_document(notes_html, document)

    document.save(docx_buffer)
    docx_buffer.seek(0)
    docx_b64 = base64.b64encode(docx_buffer.read()).decode()
    docx_name = f"{notebook_data['title'].replace(' ', '_')}.docx"

    # Build PDF bytes
    html_string = _build_html(notebook_data)
    pdf_bytes = HTML(string=html_string).write_pdf()
    pdf_b64 = base64.b64encode(pdf_bytes).decode()
    pdf_name = f"{notebook_data['title'].replace(' ', '_')}.pdf"

    # Render custom HTML with big square cards and inline SVG icons
    cards_html = f"""
    <style>
      .export-container {{
        display: flex;
        justify-content: center;
        gap: 2rem;
        margin-top: 1.5rem;
      }}
      .export-card {{
        width: 180px;
        height: 180px;
        border-radius: 16px;
        background: #1e1e1e;
        border: 1px solid rgba(255, 255, 255, 0.1);
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        text-decoration: none;
        color: white;
        font-family: sans-serif;
        box-shadow: 0 10px 25px rgba(0, 0, 0, 0.35);
        transition: transform 0.15s ease, box-shadow 0.15s ease,
                    border-color 0.15s ease, background 0.15s ease;
      }}
      .export-card:hover {{
        transform: translateY(-4px);
        box-shadow: 0 18px 35px rgba(0, 0, 0, 0.45);
        border-color: rgba(255, 255, 255, 0.3);
        background: #242424;
      }}
      .export-card svg {{
        width: 56px;
        height: 56px;
        margin-bottom: 0.75rem;
      }}
      .export-card span {{
        font-size: 0.95rem;
        font-weight: 600;
        letter-spacing: 0.03em;
      }}
    </style>

    <div class="export-container">
      <a
        class="export-card docx"
        href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}"
        download="{docx_name}"
      >
        {DOCX_SVG}
        <span>Export as DOCX</span>
      </a>
      <a
        class="export-card pdf"
        href="data:application/pdf;base64,{pdf_b64}"
        download="{pdf_name}"
      >
        {PDF_SVG}
        <span>Export as PDF</span>
      </a>
    </div>
    """

    st.components.v1.html(cards_html, height=260)


def export(notebook_data) -> None:
    """Entry point used from the main app to open the export dialog."""
    _export_dialog(notebook_data)
